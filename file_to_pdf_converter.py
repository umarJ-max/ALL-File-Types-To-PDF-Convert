import os
import sys
from pathlib import Path
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.utils import ImageReader
import docx
import openpyxl
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

class FileToPDFConverter:
    def __init__(self):
        self.supported_formats = {
            'images': ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif'],
            'text': ['.txt', '.md'],
            'word': ['.docx'],
            'excel': ['.xlsx', '.xls'],
        }
    
    def convert_to_pdf(self, input_path, output_path=None):
        """Convert file to PDF based on its extension"""
        input_path = Path(input_path)
        
        if not input_path.exists():
            raise FileNotFoundError(f"File not found: {input_path}")
        
        if output_path is None:
            output_path = input_path.with_suffix('.pdf')
        
        file_ext = input_path.suffix.lower()
        
        if file_ext in self.supported_formats['images']:
            return self._convert_image_to_pdf(input_path, output_path)
        elif file_ext in self.supported_formats['text']:
            return self._convert_text_to_pdf(input_path, output_path)
        elif file_ext in self.supported_formats['word']:
            return self._convert_word_to_pdf(input_path, output_path)
        elif file_ext in self.supported_formats['excel']:
            return self._convert_excel_to_pdf(input_path, output_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _convert_image_to_pdf(self, input_path, output_path):
        """Convert image to PDF"""
        with Image.open(input_path) as img:
            if img.mode != 'RGB':
                img = img.convert('RGB')
            img.save(output_path, 'PDF')
        return output_path
    
    def _convert_text_to_pdf(self, input_path, output_path):
        """Convert text file to PDF"""
        doc = SimpleDocTemplate(str(output_path), pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        with open(input_path, 'r', encoding='utf-8') as file:
            content = file.read()
            for line in content.split('\n'):
                if line.strip():
                    story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 12))
        
        doc.build(story)
        return output_path
    
    def _convert_word_to_pdf(self, input_path, output_path):
        """Convert Word document to PDF using pure Python (no LibreOffice/Word needed)"""
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch

        doc_in = docx.Document(str(input_path))
        pdf = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            leftMargin=inch,
            rightMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
        )

        styles = getSampleStyleSheet()

        # Custom styles
        h1 = ParagraphStyle('DocH1', parent=styles['Heading1'], fontSize=18, spaceAfter=10)
        h2 = ParagraphStyle('DocH2', parent=styles['Heading2'], fontSize=14, spaceAfter=8)
        h3 = ParagraphStyle('DocH3', parent=styles['Heading3'], fontSize=12, spaceAfter=6)
        body = ParagraphStyle('DocBody', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=6)
        bold_style = ParagraphStyle('DocBold', parent=body, fontName='Helvetica-Bold')

        heading_map = {
            'heading 1': h1,
            'heading 2': h2,
            'heading 3': h3,
        }

        def para_to_reportlab(p):
            """Convert a docx paragraph to a reportlab Paragraph, preserving basic inline formatting."""
            style_name = p.style.name.lower() if p.style and p.style.name else ''
            rl_style = heading_map.get(style_name, body)

            # Build inline-formatted text
            parts = []
            for run in p.runs:
                text = run.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                if not text:
                    continue
                if run.bold and run.italic:
                    parts.append(f'<b><i>{text}</i></b>')
                elif run.bold:
                    parts.append(f'<b>{text}</b>')
                elif run.italic:
                    parts.append(f'<i>{text}</i>')
                elif run.underline:
                    parts.append(f'<u>{text}</u>')
                else:
                    parts.append(text)

            full_text = ''.join(parts)
            if not full_text.strip():
                return None
            return Paragraph(full_text, rl_style)

        story = []

        for block in doc_in.element.body:
            tag = block.tag.split('}')[-1] if '}' in block.tag else block.tag

            if tag == 'p':
                # It's a paragraph
                from docx.oxml.ns import qn
                p_obj = None
                for p in doc_in.paragraphs:
                    if p._element is block:
                        p_obj = p
                        break
                if p_obj is not None:
                    result = para_to_reportlab(p_obj)
                    if result:
                        story.append(result)
                    else:
                        story.append(Spacer(1, 6))

            elif tag == 'tbl':
                # It's a table
                for tbl in doc_in.tables:
                    if tbl._element is block:
                        data = []
                        for row in tbl.rows:
                            row_data = []
                            for cell in row.cells:
                                cell_text = cell.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                row_data.append(Paragraph(cell_text, body))
                            data.append(row_data)

                        if data:
                            col_count = max(len(r) for r in data)
                            col_width = (A4[0] - 2 * inch) / col_count if col_count else inch

                            tbl_style = TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f0ece4')),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#0a0a0f')),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, -1), 9),
                                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#faf8f4')]),
                                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dddddd')),
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                                ('TOPPADDING', (0, 0), (-1, -1), 4),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                            ])

                            pdf_table = Table(data, colWidths=[col_width] * col_count, repeatRows=1)
                            pdf_table.setStyle(tbl_style)
                            story.append(Spacer(1, 8))
                            story.append(pdf_table)
                            story.append(Spacer(1, 8))
                        break

        if not story:
            story.append(Paragraph('(Empty document)', body))

        pdf.build(story)
        return output_path
    
    def _convert_excel_to_pdf(self, input_path, output_path):
        """Convert Excel file to PDF"""
        wb = openpyxl.load_workbook(input_path)
        doc = SimpleDocTemplate(str(output_path), pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            story.append(Paragraph(f"Sheet: {sheet_name}", styles['Heading1']))
            
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                if row_text.strip():
                    story.append(Paragraph(row_text, styles['Normal']))
            story.append(Spacer(1, 20))
        
        doc.build(story)
        return output_path
    
    def batch_convert(self, folder_path, output_folder=None):
        """Convert all supported files in a folder"""
        folder_path = Path(folder_path)
        if output_folder is None:
            output_folder = folder_path / "converted_pdfs"
        
        output_folder = Path(output_folder)
        output_folder.mkdir(exist_ok=True)
        
        converted_files = []
        all_extensions = []
        for ext_list in self.supported_formats.values():
            all_extensions.extend(ext_list)
        
        for file_path in folder_path.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in all_extensions:
                try:
                    output_path = output_folder / f"{file_path.stem}.pdf"
                    self.convert_to_pdf(file_path, output_path)
                    converted_files.append((file_path, output_path))
                    print(f"Converted: {file_path.name} → {output_path.name}")
                except Exception as e:
                    print(f"Error converting {file_path.name}: {e}")
        
        return converted_files

def main():
    converter = FileToPDFConverter()
    
    if len(sys.argv) < 2:
        print("Usage:")
        print("  Single file: python file_to_pdf_converter.py <input_file> [output_file]")
        print("  Batch convert: python file_to_pdf_converter.py --batch <folder_path> [output_folder]")
        print(f"  Supported formats: {converter.supported_formats}")
        return
    
    if sys.argv[1] == "--batch":
        if len(sys.argv) < 3:
            print("Please provide folder path for batch conversion")
            return
        
        folder_path = sys.argv[2]
        output_folder = sys.argv[3] if len(sys.argv) > 3 else None
        converter.batch_convert(folder_path, output_folder)
    else:
        input_file = sys.argv[1]
        output_file = sys.argv[2] if len(sys.argv) > 2 else None
        
        try:
            result = converter.convert_to_pdf(input_file, output_file)
            print(f"Successfully converted to: {result}")
        except Exception as e:
            print(f"Error: {e}")

if __name__ == "__main__":
    main()
